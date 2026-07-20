import os
import re
import docx

def normalize_extracted_text(text: str) -> str:
    """Clean unicode artifacts, normalize spacing, and resolve letter-spaced text."""
    if not text:
        return ""
    text = text.replace("\xa0", " ").replace("\u2013", "-").replace("\u2014", " - ")
    text = text.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\ufffd", "'")
    
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        tokens = stripped.split()
        # Detect single-character letter spacing (common in pypdf bad glyph extractions)
        if len(tokens) >= 3 and sum(1 for t in tokens if len(t) == 1) / len(tokens) > 0.4:
            fixed_line = re.sub(r'(?<=\b[A-Za-z0-9@._\-])\s+(?=[A-Za-z0-9@._\-]\b)', '', stripped)
            cleaned_lines.append(fixed_line)
        else:
            cleaned_lines.append(line)
    return "\n".join(cleaned_lines)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz) or pdfplumber if available,
    falling back to pypdf. Safe from any file read exceptions.
    """
    text_content = []

    # 1. Try PyMuPDF (fitz) - best performance for multi-column layout & kerning
    try:
        try:
            import fitz
        except ImportError:
            import pymupdf as fitz
            
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_content.append(page_text)
        extracted = "\n".join(text_content)
        if extracted.strip():
            return normalize_extracted_text(extracted)
    except Exception as fitz_err:
        print(f"[Text Parser] PyMuPDF extraction warning for {file_path}: {fitz_err}")

    # 2. Try pdfplumber as secondary engine
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        extracted = "\n".join(text_content)
        if extracted.strip():
            return normalize_extracted_text(extracted)
    except Exception as plumber_err:
        print(f"[Text Parser] pdfplumber extraction warning for {file_path}: {plumber_err}")

    # 3. Fallback to pypdf
    try:
        import pypdf
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                except Exception as page_err:
                    print(f"[Text Parser] Warning: failed to extract text from PDF page: {page_err}")
        extracted = "\n".join(text_content)
        if extracted.strip():
            return normalize_extracted_text(extracted)
    except Exception as pdf_err:
        print(f"[Text Parser] Failed to read PDF file {file_path}: {pdf_err}")

    return normalize_extracted_text("\n".join(text_content))

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file, including tables, wrapped in try-except."""
    text_content = []
    try:
        doc = docx.Document(file_path)
        # Extract text from paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                text_content.append(p.text)
                
        # Extract text from tables to capture tabular resume data
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
    except Exception as docx_err:
        print(f"[Text Parser] Failed to read DOCX file {file_path}: {docx_err}")
    return normalize_extracted_text("\n".join(text_content))

def extract_text(file_path: str) -> str:
    """Detect file type and extract text content. Safe from any file read exceptions."""
    try:
        _, ext = os.path.splitext(file_path.lower())
        if ext == ".pdf":
            text = extract_text_from_pdf(file_path)
            if text.strip():
                return text
        elif ext in [".docx", ".doc"]:
            text = extract_text_from_docx(file_path)
            if text.strip():
                return text
        
        # Fallback to plain text reading
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return normalize_extracted_text(f.read())
    except Exception as e:
        print(f"[Text Parser] Critical failure extracting {file_path}: {e}")
        
    return f"Document file reference: {os.path.basename(file_path)}"

